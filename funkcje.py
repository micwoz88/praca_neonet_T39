from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from PIL import Image
import requests
import os
import re
from re import IGNORECASE, MULTILINE
import time
import datetime
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

data = datetime.date.today()
czas = datetime.datetime.now()
znacznik_czasu = f'{data.year}{data.month}{data.day}_{czas.hour}:{czas.minute}'

adres_url = ''
nazwa_sklepu = ''
#---------------------------------------------------------------------------------------
#--------------------------pojenmik na dane---------------------------------------------
#---------------------------------------------------------------------------------------
dane = []
nr_banera = 0

def dodaj_rekord(tytul,adres='',typ=2,pod_tytul='',czas_trwania='',regulamin='',czas_regul=''):
    '''
    -['tytul','pod_tytul','typ','adres',nr_banera,'czas_trwania','regulamin','czas_regul']-
    ----0----------1--------2-----3---------4----------5------------6--------------7-------

    typ - 1-rotujący, 2-zwykły
    '''
    global nr_banera
    if typ == 1:
        typ = 'Baner rotujący na stronie głównej'
    elif typ == 2:
        typ = 'Baner na stronie głównej'
    if not(adres):
        typ = ''
    nr_banera += 1
    dane.append([tytul,pod_tytul,typ,adres,nr_banera,czas_trwania,regulamin,czas_regul])

def dodaj_https(link):
    if link.startswith('https'):
        return link
    else:
        if link[0] == '/':
            return adres_url + link
        else:
            return adres_url + '/' + link

def zapisz_do_pliku(nazwa,tresc,rozszerzenie='txt'):
    name = nazwa + '.' + rozszerzenie
    file = open(name,'w',encoding='utf-8')
    file.write(str(tresc))
    file.close()

def pobierz_kod_Beauti(adres_podany):
    return BeautifulSoup(requests.get(adres_podany).text, features='html.parser')

def pobierz_kod_selenium(adres_podany,czas=3,screen=False,minimalizuj=False,wymiar_okna=[0,0],cookie_id=''):
    global nazwa_sklepu
    if screen == True:
        nazwa_sklepu = adres_podany[adres_podany.find('/www.') + 5:-3]
    browser = webdriver.Firefox()
    browser.get(adres_podany)
    if cookie_id != '':
        time.sleep(5)
        browser.find_element_by_id(cookie_id).click()
    if wymiar_okna[0] == 0 and wymiar_okna[1] == 0:
        wysokosc_strony = browser.execute_script("return document.body.scrollHeight")
        browser.set_window_size(1600,wysokosc_strony + 8000)
    else:
        browser.set_window_size(wymiar_okna[0],wymiar_okna[1])
    if minimalizuj == True:
        browser.minimize_window()
    time.sleep(czas)
    if screen == 1:
        browser.save_screenshot(f'img/{nazwa_sklepu}.png')
    kod_zrodlowy = BeautifulSoup(str(browser.page_source), features='html.parser')
    browser.quit()
    return kod_zrodlowy

def pobierz_kod_selenium_scroll(adres_podany,czas=10):
    browser = webdriver.Firefox()
    browser.get(adres_podany)
    # Get scroll height
    last_height = browser.execute_script("return document.body.scrollHeight")
    while True:
        # Scroll down to bottom
        browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        # Wait to load page
        time.sleep(czas)
        # Calculate new scroll height and compare with last scroll height
        new_height = browser.execute_script("return document.body.scrollHeight")
        if new_height == last_height:
            break
        last_height = new_height
    kod_zrodlowy = BeautifulSoup(str(browser.page_source), features='html.parser')
    browser.close()
    return kod_zrodlowy

def zapisz_dane_do_txt(nazwa_pliku):
    name = nazwa_pliku + '.txt'
    file = open(name,'w',encoding='utf-8')
    for rekord in dane:
        file.write(rekord[0] + '\n')
        if not(rekord[1] == ''):
            file.write(rekord[1] + '\n')
        if not(rekord[3] == ''):
            file.write(rekord[2] + '\n')
            file.write(rekord[3] + '\n\n')
    file.close()

def dodaj_hiperlink(paragraf,adres):
    r_id = paragraf.part.relate_to(adres, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = adres
    hyperlink.append(new_run)
    r = paragraf.add_run ()
    r._r.append (hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink

def zapisz_dane_do_docx(nazwa_pliku='dane'):
    dokument = docx.Document()
    for rekord in range(0,len(dane)):
        run_n = 0
        dokument.add_paragraph('', 'ListBullet')
    #tytul
        dokument.paragraphs[rekord].add_run(dane[rekord][0]).bold = True
        if not(dane[rekord][2] == ''):
            dokument.paragraphs[rekord].runs[run_n].add_break()
            run_n += 1
    #podtytul
        if dane[rekord][1]:
            dokument.paragraphs[rekord].add_run(dane[rekord][1]).bold = True
            dokument.paragraphs[rekord].runs[run_n].add_break()
            run_n += 1
    #typ
        if dane[rekord][2]:
            dokument.paragraphs[rekord].add_run(dane[rekord][2])
            dokument.paragraphs[rekord].runs[run_n].add_break()
            run_n += 1
    # czas i regulamim
        if not(dane[rekord][7]):
        # czas trwania
            if dane[rekord][5]:
                dokument.paragraphs[rekord].add_run(dane[rekord][5])
                dokument.paragraphs[rekord].runs[run_n].add_break()
                run_n += 1
        # regulamin
            if dane[rekord][6]:
                dodaj_hiperlink(dokument.paragraphs[rekord],dane[rekord][6])
                dokument.paragraphs[rekord].runs[run_n].add_break()
                run_n += 1
        else:
            if dane[rekord][7]:
                for info in dane[rekord][7]:
                    # czas trwania
                    if not(info[0] == ''):
                        dokument.paragraphs[rekord].add_run(info[0])
                        dokument.paragraphs[rekord].runs[run_n].add_break()
                        run_n += 1
                    # regulamin
                    if not(info[1] == ''):
                        dodaj_hiperlink(dokument.paragraphs[rekord],info[1])
                        dokument.paragraphs[rekord].runs[run_n].add_break()
                        run_n += 1
        
    #adres
        if dane[rekord][3]:
            dodaj_hiperlink(dokument.paragraphs[rekord],dane[rekord][3])
            dokument.paragraphs[rekord].runs[run_n].add_break()
            run_n += 1
    # grafika
        if os.path.exists(f'img\{dane[rekord][4]}.png'):
            dokument.paragraphs[rekord].add_run().add_picture(f'img\{dane[rekord][4]}.png', height=docx.shared.Cm(2))
            dokument.paragraphs[rekord].runs[run_n].add_break()
            os.remove(f'img\{dane[rekord][4]}.png')
    dokument.save(f'{nazwa_pliku}.docx')

def wypisz_dane():
    for rekord in dane:
        print(rekord[0])
        if not(rekord[1] == ''):
            print(rekord[1])
            print(rekord[4])
        if not(rekord[3] == ''):
            print(rekord[2])
            print(rekord[3])
            print(rekord[4])
            print()

def zapisz_grafike(ardes,rozszezenie,sciezka_nazwa):
    try:
        os.mkdir('img')
    except:
        pass
    grafika = requests.get(ardes)
    with open(f'img\{sciezka_nazwa}.{rozszezenie}','wb') as plik:
        plik.write(grafika.content)
        plik.close

def konwertuj_webp_to_png(name):
    im =Image.open(f'img/{name}.webp').convert('RGB')
    im.save(f'img/{name}.png','png')
    os.remove(f'img/{name}.webp')

def konwertuj_jpg_to_png(name):
    im =Image.open(f'img/{name}.jpg').convert('RGB')
    im.save(f'img/{name}.png','png')
    os.remove(f'img/{name}.jpg')

def konwertuj_do_png(name,roz):
    im =Image.open(f'img/{name}.{roz}').convert('RGB')
    im.save(f'img/{name}.png','png')
    os.remove(f'img/{name}.{roz}')

def rozszerzenie_img(adres_img):
    if adres_img[len(adres_img)-5] == '.':
        return adres_img[len(adres_img)-4:]
    else:
        return adres_img[len(adres_img)-3:]

def usun_img():

    for ban in range(0,nr_banera):
        try:
            for roz in ('png','jpg','webp'):
                os.remove(f'img/{ban}.{roz}')
        except:
            pass

def pobierz_info_promo_euro(strona):
    global dane
    try:
        strona = pobierz_kod_selenium(strona)
        pat_ter = r'(Promocj|Akcja|Kod ).*?((obowi|trwa).*?do.*?)?(\d\d:\d\d|\d\d\.\d\d\.\d\d\d\d(.*?\d\d:\d\d)?)'
        pat_link = r'(?<=href=")\S*?\.pdf'
        pattern_termin = re.compile(pat_ter, flags=IGNORECASE)
        pattern_reg_link = re.compile(pat_link)
        czas = pattern_termin.finditer(str(strona))

        inf_dane = []

        for informacja in czas:
            try:
                regulamin_link = pattern_reg_link.search(str(strona), informacja.end())
                link = dodaj_https(regulamin_link.group())
            except:
                link = ''

            if str(informacja.group()).find('&nbsp;') >= 0:
                informacja = re.sub(r'obowiązuje.*?trwa do&nbsp;','trwa do ',informacja.group())
            else:
                informacja = informacja.group()

            inf_dane.append([informacja, link])

        # sprawdza i usuwa dublikaty czas
        for rek in range(len(inf_dane)):
            if rek == len(inf_dane) - 1:
                break
            else:
                if inf_dane[rek][0] == inf_dane[rek + 1][0]:
                    inf_dane[rek][0] = ''
                if inf_dane[rek][0].find('obowiązuje na euro.com.pl, we wszystkich sklepach stacjonarnych oraz dla zamówień telefonicznych pod numerem 855 855 855 i '):
                    inf_dane[rek][0] = inf_dane[rek][0].replace('obowiązuje na euro.com.pl, we wszystkich sklepach stacjonarnych oraz dla zamówień telefonicznych pod numerem 855 855 855 i ', '')
                    
        dane[len(dane) - 1][7] = inf_dane

        # sprawdza i usuwa dublikaty linków
        for rek in range(len(inf_dane)):
            if rek == len(inf_dane) - 1:
                break
            else:
                if inf_dane[rek][1] == inf_dane[rek + 1][1]:
                    inf_dane[rek][1] = ''
                    
        dane[len(dane) - 1][7] = inf_dane
    except:
        pass